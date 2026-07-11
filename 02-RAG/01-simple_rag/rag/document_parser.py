"""
离线流水线 Step 1：文档解析与预处理
支持 PDF / Word / HTML / Markdown / TXT
"""
import re
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional


@dataclass
class Document:
    """解析后的文档对象"""
    content: str                          # 纯文本内容
    source: str                           # 来源文件路径
    metadata: dict = field(default_factory=dict)  # 标题、页码等元数据


class DocumentParser:
    """多格式文档解析器"""

    def parse(self, file_path: str) -> Document:
        path = Path(file_path)
        suffix = path.suffix.lower()

        parsers = {
            ".pdf":  self._parse_pdf,
            ".docx": self._parse_docx,
            ".doc":  self._parse_docx,
            ".html": self._parse_html,
            ".htm":  self._parse_html,
            ".md":   self._parse_markdown,
            ".txt":  self._parse_text,
        }

        parser_fn = parsers.get(suffix, self._parse_text)
        content, metadata = parser_fn(path)
        content = self._clean(content)

        return Document(content=content, source=str(path), metadata=metadata)

    # ──────────────────────── 各格式解析 ────────────────────────

    def _parse_pdf(self, path: Path) -> tuple[str, dict]:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append(text)
        return "\n".join(pages), {"pages": len(reader.pages)}

    def _parse_docx(self, path: Path) -> tuple[str, dict]:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs), {"paragraphs": len(paragraphs)}

    def _parse_html(self, path: Path) -> tuple[str, dict]:
        from bs4 import BeautifulSoup
        html = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "lxml")
        # 移除脚本和样式
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()
        title = soup.title.string if soup.title else ""
        return soup.get_text(separator="\n"), {"title": title}

    def _parse_markdown(self, path: Path) -> tuple[str, dict]:
        import markdown
        from bs4 import BeautifulSoup
        md_text = path.read_text(encoding="utf-8", errors="ignore")
        html = markdown.markdown(md_text)
        soup = BeautifulSoup(html, "lxml")
        return soup.get_text(separator="\n"), {"format": "markdown"}

    def _parse_text(self, path: Path) -> tuple[str, dict]:
        content = path.read_text(encoding="utf-8", errors="ignore")
        return content, {"format": "text"}

    # ──────────────────────── 文本清洗 ────────────────────────

    def _clean(self, text: str) -> str:
        """去噪：合并多余空白行、去除不可见字符"""
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'[^\S\n]+', ' ', text)   # 水平空白压缩为单空格
        text = re.sub(r'\n{3,}', '\n\n', text)  # 多个空行合并为两行
        return text.strip()
